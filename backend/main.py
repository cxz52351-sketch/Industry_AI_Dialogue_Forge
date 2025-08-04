from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any, Literal
import httpx
import json
import os
from dotenv import load_dotenv
import asyncio
import aiofiles
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Inches
import markdown
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import re

# 加载环境变量
load_dotenv()

app = FastAPI(title="Indus AI Dialogue Forge API", version="1.0.0")

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:8080", 
        "http://127.0.0.1:8080",
        "http://localhost:8081", 
        "http://127.0.0.1:8081"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 数据模型
class Message(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    messages: List[Message]
    model: str = "deepseek-chat"
    temperature: float = 0.7
    max_tokens: int = 4000
    stream: bool = True
    output_format: str = "text"  # text, pdf, docx, markdown

class ChatResponse(BaseModel):
    id: str
    object: str
    created: int
    model: str
    choices: List[Dict[str, Any]]
    usage: Dict[str, int]

# DeepSeek API配置
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
DEEPSEEK_API_BASE = "https://api.deepseek.com/v1"

if not DEEPSEEK_API_KEY:
    print("警告: 未设置DEEPSEEK_API_KEY环境变量")

# 文件存储目录
UPLOAD_DIR = "uploads"
GENERATED_DIR = "generated"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(GENERATED_DIR, exist_ok=True)

# 文件生成函数
async def generate_markdown_file(content: str, filename: str) -> str:
    """生成Markdown文件"""
    file_path = os.path.join(GENERATED_DIR, filename)
    async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
        await f.write(content)
    return file_path

async def generate_docx_file(content: str, filename: str) -> str:
    """生成DOCX文件"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('AI生成文档', 0)
    title.alignment = 1  # 居中对齐
    
    # 添加时间戳
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()  # 空行
    
    # 处理内容，支持基本的Markdown格式
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        # 标题处理
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('**') and line.endswith('**'):
            # 粗体
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('```'):
            # 代码块开始/结束，忽略
            continue
        else:
            # 普通段落
            doc.add_paragraph(line)
    
    file_path = os.path.join(GENERATED_DIR, filename)
    doc.save(file_path)
    return file_path

async def generate_pdf_file(content: str, filename: str) -> str:
    """生成PDF文件"""
    file_path = os.path.join(GENERATED_DIR, filename)
    
    # 创建PDF文档
    doc = SimpleDocTemplate(file_path, pagesize=letter,
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=18)
    
    # 获取样式
    styles = getSampleStyleSheet()
    title_style = styles['Title']
    normal_style = styles['Normal']
    heading_style = styles['Heading1']
    
    # 创建内容列表
    story = []
    
    # 添加标题
    story.append(Paragraph("AI生成文档", title_style))
    story.append(Spacer(1, 12))
    
    # 添加时间戳
    story.append(Paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', normal_style))
    story.append(Spacer(1, 12))
    
    # 处理内容
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue
            
        # 处理标题
        if line.startswith('# '):
            story.append(Paragraph(line[2:], heading_style))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:], heading_style))
        elif line.startswith('### '):
            story.append(Paragraph(line[4:], heading_style))
        else:
            # 清理特殊字符，避免ReportLab解析错误
            clean_line = re.sub(r'[<>&]', '', line)
            story.append(Paragraph(clean_line, normal_style))
        
        story.append(Spacer(1, 6))
    
    # 构建PDF
    doc.build(story)
    return file_path

async def generate_file(content: str, output_format: str) -> Optional[Dict[str, str]]:
    """根据格式生成文件"""
    if output_format == "text":
        return None
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_id = str(uuid.uuid4())[:8]
    
    try:
        if output_format == "markdown":
            filename = f"ai_response_{timestamp}_{file_id}.md"
            file_path = await generate_markdown_file(content, filename)
            return {
                "filename": filename,
                "url": f"/api/files/{filename}",
                "mime_type": "text/markdown"
            }
        elif output_format == "docx":
            filename = f"ai_response_{timestamp}_{file_id}.docx"
            file_path = await generate_docx_file(content, filename)
            return {
                "filename": filename,
                "url": f"/api/files/{filename}",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        elif output_format == "pdf":
            filename = f"ai_response_{timestamp}_{file_id}.pdf"
            file_path = await generate_pdf_file(content, filename)
            return {
                "filename": filename,
                "url": f"/api/files/{filename}",
                "mime_type": "application/pdf"
            }
    except Exception as e:
        print(f"文件生成错误: {str(e)}")
        import traceback
        print(f"错误堆栈: {traceback.format_exc()}")
        return None
    
    return None

@app.get("/")
async def root():
    return {"message": "Indus AI Dialogue Forge API", "status": "running"}

@app.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.post("/api/chat")
async def chat_with_deepseek(request: ChatRequest):
    """
    与DeepSeek模型进行对话
    """
    if not DEEPSEEK_API_KEY:
        raise HTTPException(status_code=500, detail="DeepSeek API密钥未配置")
    
    try:
        print(f"开始处理聊天请求，API Key: {DEEPSEEK_API_KEY[:8]}...")
        # 准备请求数据
        payload = {
            "model": request.model,
            "messages": [{"role": msg.role, "content": msg.content} for msg in request.messages],
            "temperature": request.temperature,
            "max_tokens": request.max_tokens,
            "stream": request.stream
        }
        
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        
        if request.stream:
            # 流式响应
            print("处理流式响应请求...")
            
            async def generate():
                full_content = ""  # 收集完整内容用于文件生成
                try:
                    async with httpx.AsyncClient(timeout=60.0) as client:
                        async with client.stream("POST", f"{DEEPSEEK_API_BASE}/chat/completions", 
                                            json=payload, headers=headers) as response:
                            if response.status_code != 200:
                                error_text = await response.aread()
                                error_msg = f"DeepSeek API错误: {error_text.decode()}"
                                print(f"API错误: {error_msg}")
                                yield f"data: {{\"error\": {{\"message\": \"{error_msg}\"}} }}\n\n"
                                yield "data: [DONE]\n\n"
                                return
                            
                            print("开始接收流式数据...")
                            buffer = ""
                            async for chunk in response.aiter_text():
                                if chunk:
                                    buffer += chunk
                                    # 处理完整的行
                                    while '\n' in buffer:
                                        line, buffer = buffer.split('\n', 1)
                                        if line.strip():
                                            if line.startswith('data: '):
                                                # 解析JSON并提取内容
                                                try:
                                                    data = line[6:].strip()
                                                    if data and data != '[DONE]':
                                                        parsed = json.loads(data)
                                                        if parsed.get('choices') and parsed['choices'][0].get('delta', {}).get('content'):
                                                            content = parsed['choices'][0]['delta']['content']
                                                            full_content += content
                                                except:
                                                    pass
                                                
                                                print(f"转发数据: {line[:50]}...")
                                                yield f"{line}\n\n"
                                            else:
                                                yield f"data: {line}\n\n"
                            
                            # 处理剩余的buffer
                            if buffer.strip():
                                if buffer.startswith('data: '):
                                    yield f"{buffer}\n\n"
                                else:
                                    yield f"data: {buffer}\n\n"
                            
                            # 生成文件（如果需要）
                            if request.output_format != "text" and full_content.strip():
                                print(f"生成{request.output_format}文件...")
                                file_info = await generate_file(full_content, request.output_format)
                                if file_info:
                                    file_data = {
                                        "type": "file",
                                        "filename": file_info["filename"],
                                        "url": file_info["url"],
                                        "mime_type": file_info["mime_type"]
                                    }
                                    yield f"data: {json.dumps(file_data)}\n\n"
                            
                            yield "data: [DONE]\n\n"
                except httpx.StreamClosed:
                    print("流连接被关闭")
                    error_response = {
                        "error": {
                            "message": "Stream was closed unexpectedly"
                        }
                    }
                    yield f"data: {json.dumps(error_response)}\n\n"
                    yield "data: [DONE]\n\n"
                except Exception as e:
                    print(f"流式处理错误: {str(e)}")
                    error_response = {
                        "error": {
                            "message": f"流式处理错误: {str(e)}"
                        }
                    }
                    yield f"data: {json.dumps(error_response)}\n\n"
                    yield "data: [DONE]\n\n"
            
            return StreamingResponse(generate(), media_type="text/event-stream")
        
        else:
            # 普通响应
            print("发送请求到 DeepSeek API...")
            print(f"请求数据: {json.dumps(payload, ensure_ascii=False, indent=2)}")
            
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(f"{DEEPSEEK_API_BASE}/chat/completions", 
                                           json=payload, headers=headers)
                
                print(f"收到响应，状态码: {response.status_code}")
                if response.status_code != 200:
                    raise HTTPException(status_code=response.status_code, 
                                      detail=f"DeepSeek API错误: {response.text}")
                
                response_data = response.json()
                print(f"响应数据: {json.dumps(response_data, ensure_ascii=False, indent=2)}")
                
                # 生成文件（如果需要）
                if request.output_format != "text" and response_data.get('choices'):
                    content = response_data['choices'][0]['message']['content']
                    if content.strip():
                        print(f"生成{request.output_format}文件...")
                        file_info = await generate_file(content, request.output_format)
                        if file_info:
                            response_data['file'] = file_info
                
                return response_data
                
    except httpx.TimeoutException:
        raise HTTPException(status_code=408, detail="请求超时")
    except httpx.RequestError as e:
        print(f"网络请求错误: {str(e)}")
        raise HTTPException(status_code=500, detail=f"网络请求错误: {str(e)}")
    except Exception as e:
        print(f"服务器内部错误: {str(e)}")
        import traceback
        print(f"错误堆栈: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"服务器内部错误: {str(e)}")

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    """
    上传文件
    """
    try:
        # 验证文件类型
        allowed_extensions = {'.pdf', '.docx', '.doc', '.txt', '.jpg', '.jpeg', '.png', '.gif', '.xlsx', '.xls', '.pptx', '.ppt'}
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension not in allowed_extensions:
            raise HTTPException(status_code=400, detail="不支持的文件类型")
        
        # 验证文件大小 (50MB)
        content = await file.read()
        file_size = len(content)
        if file_size > 50 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="文件大小超过50MB限制")
        
        # 生成唯一文件名
        file_id = str(uuid.uuid4())
        file_extension = os.path.splitext(file.filename)[1]
        new_filename = f"{file_id}{file_extension}"
        file_path = os.path.join(UPLOAD_DIR, new_filename)
        
        # 保存文件
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(content)
        
        return {
            "file_id": file_id,
            "filename": file.filename,
            "size": file_size,
            "type": file.content_type,
            "uploaded_at": datetime.now().isoformat()
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件上传失败: {str(e)}")

@app.get("/api/files/{filename}")
async def download_file(filename: str):
    """
    下载生成的文件
    """
    file_path = os.path.join(GENERATED_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 根据文件扩展名设置MIME类型
    if filename.endswith('.pdf'):
        media_type = 'application/pdf'
    elif filename.endswith('.docx'):
        media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif filename.endswith('.md'):
        media_type = 'text/markdown'
    else:
        media_type = 'application/octet-stream'
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type=media_type
    )

@app.get("/api/models")
async def get_available_models():
    """
    获取可用的模型列表
    """
    return {
        "models": [
            {
                "id": "deepseek-chat",
                "name": "DeepSeek Chat",
                "description": "DeepSeek通用对话模型",
                "max_tokens": 8192
            },
            {
                "id": "deepseek-coder",
                "name": "DeepSeek Coder", 
                "description": "DeepSeek代码生成模型",
                "max_tokens": 8192
            }
        ]
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)